#!/usr/bin/env python3
"""
generate.py — Pipeline completo (single source of truth: content.json)

  1. Valida content.json
  2. Copia para website/src/content.json  (website React)
  3. Gera cv-documents/CV_Filipe_Fernandes_EN.docx + PT.docx
  4. Converte para cv-documents/CV_Filipe_Fernandes_EN.pdf  + PT.pdf

Uso: python3 generate.py
"""

import json, shutil, subprocess, sys, re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ModuleNotFoundError:
    print("ERRO: python-docx não está instalado.")
    print("Instala com:  pip install python-docx --break-system-packages")
    sys.exit(1)

BASE     = Path(__file__).parent
SRC      = BASE / 'content.json'
DEST     = BASE / 'website' / 'src' / 'content.json'
OUT_DIR  = BASE / 'cv-documents'

# ── Colours ───────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x1F, 0x4E, 0x79)
BLUE  = RGBColor(0x2E, 0x75, 0xB6)
DARK  = RGBColor(0x1A, 0x1A, 0x1A)
GRAY  = RGBColor(0x55, 0x55, 0x55)
LIGHT = RGBColor(0x88, 0x88, 0x88)

# ── Static content (não muda com a língua) ────────────────────────────────────
STATIC = {
    'en': {
        'headline':   'System Administrator  ·  Infrastructure & Security  ·  Automation',
        'sum_label':  'Professional Summary',
        'exp_label':  'Professional Experience',
        'sk_label':   'Technical Skills',
        'dev_label':  'Professional Development',
        'edu_label':  'Education',
        'lang_label': 'Languages',
        'dev': [
            ('Continuous Learning',
             '— Networking, Fortinet, and Security courses (Udemy / Pluralsight / Coursera)'),
            ('Personal Homelab',
             '— Proxmox VE hypervisor hosting Linux and Windows VMs; Docker-based self-hosted services '
             'for testing infrastructure patterns, automation playbooks, and monitoring stacks in a '
             'production-like environment'),
        ],
        'edu':   ('Técnico de Programação e Gestão de Sistemas Informáticos',
                  '2014 – 2017',
                  'Escola Secundária Augusto Cabrita, Barreiro'),
        'langs': [('Portuguese:', 'Native'), ('English:', 'Professional Working Proficiency')],
    },
    'pt': {
        'headline':   'Administrador de Sistemas  ·  Infraestrutura & Segurança  ·  Automação',
        'sum_label':  'Resumo Profissional',
        'exp_label':  'Experiência Profissional',
        'sk_label':   'Competências Técnicas',
        'dev_label':  'Desenvolvimento Profissional',
        'edu_label':  'Formação',
        'lang_label': 'Idiomas',
        'dev': [
            ('Aprendizagem Contínua',
             '— Cursos de Redes, Fortinet e Segurança (Udemy / Pluralsight / Coursera)'),
            ('Homelab Pessoal',
             '— Hypervisor Proxmox VE com VMs Linux e Windows; serviços self-hosted em Docker '
             'para teste de padrões de infraestrutura, playbooks de automação e stacks de '
             'monitorização em ambiente semelhante ao de produção'),
        ],
        'edu':   ('Técnico de Programação e Gestão de Sistemas Informáticos',
                  '2014 – 2017',
                  'Escola Secundária Augusto Cabrita, Barreiro'),
        'langs': [('Português:', 'Nativo'), ('Inglês:', 'Proficiência Profissional de Trabalho')],
    },
}

# ── Low-level XML helpers ─────────────────────────────────────────────────────

def _hex(color: RGBColor) -> str:
    return f'{color[0]:02X}{color[1]:02X}{color[2]:02X}'

def _set_color(run, color: RGBColor):
    run.font.color.rgb = color

def _bottom_border(para, color: RGBColor, sz: int = 8):
    """Add a bottom border to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '6')
    bottom.set(qn('w:color'), _hex(color))
    pBdr.append(bottom)
    pPr.append(pBdr)

def _spacing(para, before: int = 0, after: int = 0):
    pPr = para._p.get_or_add_pPr()
    pPrSpacing = pPr.find(qn('w:spacing'))
    if pPrSpacing is None:
        pPrSpacing = OxmlElement('w:spacing')
        pPr.append(pPrSpacing)
    if before:
        pPrSpacing.set(qn('w:before'), str(before))
    if after:
        pPrSpacing.set(qn('w:after'), str(after))

def _add_tab_stop_right(para):
    """Right-align tab stop at page margin."""
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9026')   # ~16cm — right margin
    tabs.append(tab)
    pPr.append(tabs)

def _bullet_para(doc, text: str):
    """Add a bullet paragraph using the list style."""
    para = doc.add_paragraph(style='List Paragraph')
    # numbering (bullet) via XML
    pPr  = para._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl  = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), '0')
    numId = OxmlElement('w:numId'); numId.set(qn('w:val'), '1')
    numPr.append(ilvl); numPr.append(numId)
    pPr.append(numPr)

    run = para.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(9.5)
    _set_color(run, DARK)
    _spacing(para, after=30)
    return para

def _add_numbering(doc):
    """Inject a bullet numbering definition into the document."""
    numbering_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
  xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="bullet"/>
      <w:lvlText w:val="&#x2022;"/>
      <w:lvlJc w:val="left"/>
      <w:pPr>
        <w:ind w:left="480" w:hanging="240"/>
      </w:pPr>
      <w:rPr>
        <w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/>
      </w:rPr>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>'''
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    import lxml.etree as etree

    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        # create from scratch
        part = Part(
            PackURI('/word/numbering.xml'),
            'application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml',
            etree.fromstring(numbering_xml.encode()),
            doc.part.package,
        )
        doc.part.relate_to(part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering')

# ── Section helpers ───────────────────────────────────────────────────────────

def add_section_header(doc, text: str):
    para = doc.add_paragraph()
    run  = para.add_run(text.upper())
    run.bold       = True
    run.font.name  = 'Calibri'
    run.font.size  = Pt(10)
    _set_color(run, NAVY)
    _bottom_border(para, NAVY)
    _spacing(para, before=200, after=100)
    return para

def add_role(doc, company, title, period, location):
    # Line 1: Company \t Period
    p1 = doc.add_paragraph()
    r_company = p1.add_run(company)
    r_company.bold = True; r_company.font.name = 'Calibri'; r_company.font.size = Pt(10.5)
    _set_color(r_company, DARK)
    r_tab = p1.add_run('\t')
    r_period = p1.add_run(period)
    r_period.font.name = 'Calibri'; r_period.font.size = Pt(9)
    _set_color(r_period, LIGHT)
    _add_tab_stop_right(p1)
    _spacing(p1, before=140, after=20)

    # Line 2: Title \t Location
    p2 = doc.add_paragraph()
    r_title = p2.add_run(title)
    r_title.bold = True; r_title.font.name = 'Calibri'; r_title.font.size = Pt(10)
    _set_color(r_title, BLUE)
    p2.add_run('\t')
    r_loc = p2.add_run(location)
    r_loc.font.name = 'Calibri'; r_loc.font.size = Pt(9); r_loc.italic = True
    _set_color(r_loc, LIGHT)
    _add_tab_stop_right(p2)
    _spacing(p2, before=0, after=50)

def add_skill_line(doc, category, items_list):
    para = doc.add_paragraph()
    r_cat = para.add_run(category + ':  ')
    r_cat.bold = True; r_cat.font.name = 'Calibri'; r_cat.font.size = Pt(9.5)
    _set_color(r_cat, NAVY)
    r_items = para.add_run('  ·  '.join(items_list))
    r_items.font.name = 'Calibri'; r_items.font.size = Pt(9.5)
    _set_color(r_items, DARK)
    _spacing(para, after=44)

# ── Main document builder ─────────────────────────────────────────────────────

def build_cv(lang: str, content: dict) -> Document:
    c  = content[lang]
    st = STATIC[lang]

    doc = Document()

    # Page size A4 + margins
    section = doc.sections[0]
    section.page_width  = Twips(11906)
    section.page_height = Twips(16838)
    section.top_margin    = Twips(900)
    section.bottom_margin = Twips(900)
    section.left_margin   = Twips(1000)
    section.right_margin  = Twips(1000)

    # Default paragraph spacing = 0
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].paragraph_format.space_after  = Pt(0)

    _add_numbering(doc)

    # ── HEADER ────────────────────────────────────────────────────────────────
    p_name = doc.add_paragraph()
    r_name = p_name.add_run('Filipe Fernandes')
    r_name.bold = True; r_name.font.name = 'Calibri'; r_name.font.size = Pt(28)
    _set_color(r_name, NAVY)
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p_name, after=60)

    p_hl = doc.add_paragraph()
    r_hl = p_hl.add_run(st['headline'])
    r_hl.font.name = 'Calibri'; r_hl.font.size = Pt(11)
    _set_color(r_hl, GRAY)
    p_hl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p_hl, after=80)

    p_contact = doc.add_paragraph()
    for run_text, color in [
        ('+351 961 500 099  ·  ', GRAY),
        ('filipe.fernandes.work@gmail.com', BLUE),
        ('  ·  ', GRAY),
        ('linkedin.com/in/ffernandes97', BLUE),
        ('  ·  Barreiro, Portugal', GRAY),
    ]:
        r = p_contact.add_run(run_text)
        r.font.name = 'Calibri'; r.font.size = Pt(9)
        _set_color(r, color)
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _bottom_border(p_contact, NAVY, sz=6)
    _spacing(p_contact, after=30)

    # ── SUMMARY ───────────────────────────────────────────────────────────────
    add_section_header(doc, st['sum_label'])
    p_sum = doc.add_paragraph()
    r_sum = p_sum.add_run(c['about']['cv_summary'])
    r_sum.font.name = 'Calibri'; r_sum.font.size = Pt(9.5)
    _set_color(r_sum, DARK)
    p_sum.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _spacing(p_sum, before=60, after=60)

    # ── EXPERIENCE ────────────────────────────────────────────────────────────
    add_section_header(doc, st['exp_label'])
    for job in c['experience']['jobs']:
        add_role(doc, job['company'], job['role'], job['period'], job['location'])
        for bullet in job['bullets']:
            _bullet_para(doc, bullet)

    # ── SKILLS ────────────────────────────────────────────────────────────────
    add_section_header(doc, st['sk_label'])
    for group in c['skills']['groups']:
        add_skill_line(doc, group['category'], group['items'])

    # ── PROFESSIONAL DEVELOPMENT ──────────────────────────────────────────────
    add_section_header(doc, st['dev_label'])
    for bold_text, rest_text in st['dev']:
        p = doc.add_paragraph()
        r_b = p.add_run(bold_text + '  ')
        r_b.bold = True; r_b.font.name = 'Calibri'; r_b.font.size = Pt(9.5)
        _set_color(r_b, NAVY)
        r_r = p.add_run(rest_text)
        r_r.font.name = 'Calibri'; r_r.font.size = Pt(9.5)
        _set_color(r_r, DARK)
        _spacing(p, before=60, after=50)

    # ── EDUCATION ─────────────────────────────────────────────────────────────
    add_section_header(doc, st['edu_label'])
    degree, period, school = st['edu']
    p_deg = doc.add_paragraph()
    r_deg = p_deg.add_run(degree)
    r_deg.bold = True; r_deg.font.name = 'Calibri'; r_deg.font.size = Pt(10)
    _set_color(r_deg, DARK)
    p_deg.add_run('\t')
    r_per = p_deg.add_run(period)
    r_per.font.name = 'Calibri'; r_per.font.size = Pt(9)
    _set_color(r_per, LIGHT)
    _add_tab_stop_right(p_deg)
    _spacing(p_deg, before=60, after=20)

    p_sch = doc.add_paragraph()
    r_sch = p_sch.add_run(school)
    r_sch.font.name = 'Calibri'; r_sch.font.size = Pt(9); r_sch.italic = True
    _set_color(r_sch, GRAY)
    _spacing(p_sch, after=50)

    # ── LANGUAGES ─────────────────────────────────────────────────────────────
    add_section_header(doc, st['lang_label'])
    p_lang = doc.add_paragraph()
    for i, (label, value) in enumerate(st['langs']):
        r_l = p_lang.add_run(label + ' ')
        r_l.bold = True; r_l.font.name = 'Calibri'; r_l.font.size = Pt(9.5)
        _set_color(r_l, DARK)
        r_v = p_lang.add_run(value)
        r_v.font.name = 'Calibri'; r_v.font.size = Pt(9.5)
        _set_color(r_v, DARK)
        if i < len(st['langs']) - 1:
            r_sep = p_lang.add_run('   ·   ')
            r_sep.font.name = 'Calibri'; r_sep.font.size = Pt(9.5)
            _set_color(r_sep, LIGHT)
    _spacing(p_lang, before=60, after=40)

    return doc

# ── Steps ─────────────────────────────────────────────────────────────────────

def step_validate(content_path: Path) -> dict:
    try:
        with open(content_path, encoding='utf-8') as f:
            data = json.load(f)
    except json.JSONDecodeError as e:
        print(f'ERROR: content.json is invalid JSON — {e}'); sys.exit(1)

    SECTIONS = ['nav', 'home', 'about', 'experience', 'skills', 'projects', 'certifications', 'contact']
    ok = True
    for lang in ('en', 'pt'):
        if lang not in data:
            print(f'ERROR: missing language "{lang}"'); ok = False; continue
        for s in SECTIONS:
            if s not in data[lang]:
                print(f'ERROR: missing "{lang}.{s}"'); ok = False
        if 'cv_summary' not in data[lang].get('about', {}):
            print(f'ERROR: missing "{lang}.about.cv_summary" — needed for DOCX generation'); ok = False

    if not ok:
        sys.exit(1)

    exp_en  = len(data['en']['experience']['jobs'])
    proj_en = len(data['en']['projects']['list'])
    print('✓ content.json valid')
    print(f'  EN: {exp_en} jobs, {proj_en} projects')
    print(f'  PT: {len(data["pt"]["experience"]["jobs"])} jobs, {len(data["pt"]["projects"]["list"])} projects')
    return data


def step_sync_website(src: Path, dest: Path):
    shutil.copy2(src, dest)
    print('✓ Copied → website/src/content.json')


def step_generate_docx(data: dict, out_dir: Path):
    out_dir.mkdir(exist_ok=True)
    for lang in ('en', 'pt'):
        out = out_dir / f'CV_Filipe_Fernandes_{lang.upper()}.docx'
        doc = build_cv(lang, data)
        doc.save(out)
        print(f'✓ {out.name}')


def _soffice_bin():
    """Encontra soffice/libreoffice no PATH."""
    for name in ('soffice', 'libreoffice'):
        r = subprocess.run(['which', name], capture_output=True, text=True)
        if r.returncode == 0:
            return r.stdout.strip()
    return None


def _wsl_to_win(path: Path) -> str:
    """Converte path WSL para Windows (ex: /mnt/c/... → C:\\...)."""
    r = subprocess.run(['wslpath', '-w', str(path)], capture_output=True, text=True)
    return r.stdout.strip() if r.returncode == 0 else str(path)


def _convert_via_word(docx_path: Path, out_dir: Path) -> bool:
    """Usa Word COM via PowerShell (fallback WSL/Windows)."""
    win_docx = _wsl_to_win(docx_path)
    win_out  = _wsl_to_win(out_dir)
    ps_script = (
        "$w = New-Object -ComObject Word.Application; "
        "$w.Visible = $false; "
        f"$d = $w.Documents.Open('{win_docx}'); "
        f"$d.SaveAs2('{win_out}\\\\{docx_path.stem}.pdf', 17); "
        "$d.Close(); $w.Quit()"
    )
    r = subprocess.run(
        ['powershell.exe', '-NoProfile', '-Command', ps_script],
        capture_output=True, text=True
    )
    return r.returncode == 0


def step_convert_pdf(out_dir: Path):
    soffice = _soffice_bin()
    for lang in ('EN', 'PT'):
        docx_path = out_dir / f'CV_Filipe_Fernandes_{lang}.docx'
        pdf_path  = out_dir / f'CV_Filipe_Fernandes_{lang}.pdf'

        if soffice:
            # ── Opção 1: LibreOffice ──
            r = subprocess.run(
                [soffice, '--headless', '--convert-to', 'pdf',
                 str(docx_path), '--outdir', str(out_dir)],
                capture_output=True, text=True
            )
            ok = r.returncode == 0
        else:
            # ── Opção 2: Word via PowerShell (WSL/Windows) ──
            ok = _convert_via_word(docx_path, out_dir)

        if ok and pdf_path.exists():
            print(f'✓ CV_Filipe_Fernandes_{lang}.pdf')
        else:
            print(f'⚠ PDF não gerado para {lang} — instala LibreOffice: sudo apt install libreoffice')
            print(f'  DOCX disponível em: {docx_path}')


def step_sync_public(cv_dir: Path, public_dir: Path):
    """Keep website/public/ in sync with the generated PDFs."""
    public_dir.mkdir(exist_ok=True)
    for lang in ('EN', 'PT'):
        src = cv_dir / f'CV_Filipe_Fernandes_{lang}.pdf'
        dst = public_dir / src.name
        shutil.copy2(src, dst)
    print('✓ PDFs copied → website/public/')

# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == '__main__':
    print('\n── generate.py ─────────────────────────────────────')

    data = step_validate(SRC)
    step_sync_website(SRC, DEST)
    step_generate_docx(data, OUT_DIR)
    step_convert_pdf(OUT_DIR)
    step_sync_public(OUT_DIR, BASE / 'website' / 'public')

    print('────────────────────────────────────────────────────')
    print('Done. Edit content.json → python3 generate.py → tudo sincronizado.\n')
